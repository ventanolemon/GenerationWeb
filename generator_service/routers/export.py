"""
POST /export — экспорт заданий в .docx.

Возвращает бинарный StreamingResponse — ASP.NET Core потом пробросит
его в Results.File для скачивания браузером.

Под капотом — ровно тот же стек, что и в десктопе: docx.Document
плюс block.render_docx(). Это работает в headless-окружении:
  - python-docx чистый Python без Qt;
  - matplotlib используется для рендера формул (через latex_to_png_bytes),
    у него уже стоит backend "Agg" в core/rendering.py;
  - ImageBlock.render_docx использует PIL и BytesIO.

Экспорт идёт только для StaticTask. Если генератор интерактивный,
возвращаем 400 — фронт должен это знать заранее по флагу capabilities
у раздела (этим займётся ASP.NET-слой) и не показывать кнопку.
"""

from __future__ import annotations
from io import BytesIO

from docx import Document
from fastapi import APIRouter, HTTPException, Request
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field

from core import StaticTask

router = APIRouter(prefix="/export", tags=["export"])


class ExportRequest(BaseModel):
    partition_id: int = Field(..., gt=0)
    count: int = Field(1, ge=1, le=50, description="Сколько заданий сгенерировать")
    with_answers: bool = Field(True)


@router.post("")
def export_tasks(body: ExportRequest, request: Request):
    registry = request.app.state.registry
    repo = request.app.state.repo

    if not registry.has(body.partition_id):
        raise HTTPException(
            status_code=404,
            detail=f"Generator not found for partition {body.partition_id}",
        )

    partition = repo.get_partition(body.partition_id)
    params = partition.generation_params if partition else {}
    generator = registry.get(body.partition_id, params)

    doc = Document()
    title = partition.name if partition else "Задания"
    doc.add_heading(title, level=0)

    for i in range(body.count):
        task = generator.generate()
        if not isinstance(task, StaticTask):
            raise HTTPException(
                status_code=400,
                detail="Export is only available for static tasks",
            )
        doc.add_heading(f"Задание {i + 1}", level=2)
        for block in task.statement:
            block.render_docx(doc)
        if body.with_answers:
            doc.add_heading(f"Ответ {i + 1}", level=3)
            for block in task.answer:
                block.render_docx(doc)
        if i < body.count - 1:
            doc.add_page_break()

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = f"tasks_{body.partition_id}.docx"
    return StreamingResponse(
        buf,
        media_type=(
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        ),
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
